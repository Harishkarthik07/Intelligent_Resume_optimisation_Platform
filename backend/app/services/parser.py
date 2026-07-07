import io
import re
import logging

logger = logging.getLogger(__name__)


class ResumeParser:
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }

    def parse(self, file_bytes: bytes, content_type: str) -> str:
        file_type = self.SUPPORTED_TYPES.get(content_type)
        if not file_type:
            raise ValueError(f"Unsupported file type: {content_type}. Use PDF, DOCX, or TXT.")

        logger.info(f"Parsing {file_type} file ({len(file_bytes)} bytes)")

        if file_type == "pdf":
            return self._parse_pdf(file_bytes)
        elif file_type == "docx":
            return self._parse_docx(file_bytes)
        else:
            return self._clean(file_bytes.decode("utf-8", errors="ignore"))

    def _parse_pdf(self, data: bytes) -> str:
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(
                        x_tolerance=3,
                        y_tolerance=3,
                        layout=True,
                    )
                    if page_text:
                        text_parts.append(page_text)
            return self._clean("\n".join(text_parts))
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying pypdf")
            return self._parse_pdf_fallback(data)

    def _parse_pdf_fallback(self, data: bytes) -> str:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
        return self._clean("\n".join(parts))

    def _parse_docx(self, data: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return self._clean("\n".join(paragraphs))

    @staticmethod
    def _clean(text: str) -> str:
        text = re.sub(r'\r\n|\r', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'[^\x00-\x7F\n]', ' ', text)
        return text.strip()


resume_parser = ResumeParser()
